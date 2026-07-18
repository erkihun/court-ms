from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Erkis\Documents\GitHub\court-ms")
SOURCE = ROOT / "CCMS_User_Manual_EN_AM_v1.0.docx"
OUT_DIR = ROOT / "output" / "doc"
OUTPUT = OUT_DIR / "CCMS_User_Manual_EN_AM_v1.1_Detailed.docx"

BLUE = "17365D"
MID_BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
PALE = "F3F6F9"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name, size=9, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def set_cell_text(cell, lines, amharic=False, bold_first=False):
    cell.text = ""
    font = "Nyala" if amharic else "Arial"
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.04
        r = p.add_run(line)
        set_run_font(r, font, 9.2 if amharic else 8.7, bold=(bold_first and idx == 0))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    margins(cell)


DETAILS = {
5: (["Before you begin: Use your own active email address and have your identity/contact information ready.", "Procedure: Complete every required field, read the terms, select the required confirmation, and submit once.", "Expected result: The system creates a pending account and sends an email verification instruction.", "Control: Do not create duplicate accounts or share credentials."], ["ከመጀመርዎ በፊት፡ የራስዎን የሚሠራ ኢሜይል ይጠቀሙ፤ የማንነትና የመገኛ መረጃዎን ያዘጋጁ።", "የአፈጻጸም ሂደት፡ ግዴታ መስኮችን ይሙሉ፣ ደንቦቹን ያንብቡ፣ ማረጋገጫውን ይምረጡና አንድ ጊዜ ያስገቡ።", "የሚጠበቅ ውጤት፡ ሥርዓቱ በማረጋገጥ ላይ ያለ መለያ ፈጥሮ የኢሜይል ማረጋገጫ መመሪያ ይልካል።", "ቁጥጥር፡ ተደጋጋሚ መለያ አይፍጠሩ፤ የመግቢያ መረጃዎን ለሌላ ሰው አያጋሩ።"]),
6: (["Security: Enter the OTP only on the official CCMS page. Never send it by message or telephone.", "If it expires or is not received, use Resend once, check spam/junk, and confirm that the registered email is correct.", "Expected result: A verified account can sign in; an unverified account remains restricted."], ["ደህንነት፡ የአንድ ጊዜ ኮዱን በይፋዊው CCMS ገጽ ላይ ብቻ ያስገቡ፤ በመልዕክት ወይም በስልክ ለማንም አይስጡ።", "ኮዱ ጊዜው ካለፈ ወይም ካልደረሰ Resend ን አንድ ጊዜ ይጠቀሙ፣ spam/junk ይፈትሹ፣ የተመዘገበው ኢሜይል ትክክል መሆኑን ያረጋግጡ።", "የሚጠበቅ ውጤት፡ የተረጋገጠ መለያ ወደ ሥርዓቱ ይገባል፤ ያልተረጋገጠ መለያ ገደብ ይኖረዋል።"]),
7: (["Use Forgot password only for your registered email. Follow the reset link promptly and create a new, unique password.", "After repeated failure, stop retrying, verify keyboard language and Caps Lock, then contact authorized support.", "Always sign out on a shared or public computer."], ["Forgot password የሚለውን ለተመዘገበው ኢሜይልዎ ብቻ ይጠቀሙ። የመልሶ ማቋቋሚያ አገናኙን በወቅቱ ተከትለው አዲስና ልዩ የይለፍ ቃል ይፍጠሩ።", "በተደጋጋሚ ካልተሳካ ሙከራውን ያቁሙ፣ የቁልፍ ሰሌዳ ቋንቋና Caps Lock ይፈትሹ፣ ከዚያም የተፈቀደለትን ድጋፍ ያነጋግሩ።", "በጋራ ወይም በሕዝብ ኮምፒውተር ላይ ሁልጊዜ ከመለያዎ ይውጡ።"]),
8: (["Preparation: Gather the parties' legal names, addresses, subject, facts, requested remedy, and supporting records.", "Enter facts clearly and chronologically. Save a draft when available and review names, dates, amounts, and jurisdiction before continuing.", "Only submit information relevant to the case. False, duplicate, or incomplete data may delay institutional review."], ["ዝግጅት፡ የተከራካሪ ወገኖችን ሕጋዊ ስም፣ አድራሻ፣ የጉዳዩን ርዕስ፣ ፍሬ ነገር፣ የተጠየቀውን ውሳኔና ደጋፊ መዝገቦች ያዘጋጁ።", "ፍሬ ነገሩን በግልጽና በጊዜ ቅደም ተከተል ያስገቡ። ረቂቅ ማስቀመጥ ካለ ይጠቀሙ፤ ስሞችን፣ ቀኖችን፣ መጠኖችንና ሥልጣን ይገምግሙ።", "ከጉዳዩ ጋር የሚዛመድ መረጃ ብቻ ያስገቡ። ሐሰተኛ፣ ተደጋጋሚ ወይም ያልተሟላ መረጃ የተቋሙን ግምገማ ሊያዘገይ ይችላል።"]),
9: (["Give files meaningful names and verify that each opens correctly before upload. Follow the type and size rules displayed by the system.", "For each witness, record accurate identity and contact details and summarize what the witness can establish.", "Sensitive personal data must be limited to what the proceeding lawfully requires."], ["ፋይሎችን ገላጭ በሆነ ስም ይሰይሙ፤ ከመጫንዎ በፊት እያንዳንዱ በትክክል መከፈቱን ያረጋግጡ። ሥርዓቱ የሚያሳየውን የፋይል ዓይነትና መጠን ገደብ ይከተሉ።", "ለእያንዳንዱ ምስክር ትክክለኛ የማንነትና የመገኛ መረጃ ያስገቡ፤ ምስክሩ ሊያረጋግጥ የሚችለውን ፍሬ ነገር በአጭሩ ይግለጹ።", "ሚስጥራዊ የግል መረጃ ለሕጋዊው ሂደት በሚያስፈልገው መጠን ብቻ መቅረብ አለበት።"]),
10: (["Final check: Confirm parties, claims, attachments, witnesses, declarations, and contact information.", "After submission, record the displayed case number and download or print the receipt. Keep both in a secure place.", "Submission does not by itself mean acceptance. Monitor the review status and respond to institutional requests."], ["የመጨረሻ ማረጋገጫ፡ ተከራካሪ ወገኖችን፣ ጥያቄዎችን፣ አባሪዎችን፣ ምስክሮችን፣ መግለጫዎችንና የመገኛ መረጃን ያረጋግጡ።", "ካስገቡ በኋላ የታየውን የጉዳይ መዝገብ ቁጥር ይመዝግቡ፤ ደረሰኙን ያውርዱ ወይም ያትሙ። ሁለቱንም በደህና ያስቀምጡ።", "ማስገባት ብቻውን ተቀባይነትን አያመለክትም። የግምገማ ሁኔታውን ይከታተሉ፤ ለተቋሙ ጥያቄዎች በወቅቱ ምላሽ ይስጡ።"]),
11: (["Open the case detail rather than relying only on the dashboard label. Review status history, hearings, decisions, files, and action requests.", "A status change is authoritative only when displayed in the system or issued through an approved institutional notice.", "Report an apparent error with the case number, screenshot, date/time, and a concise description."], ["በዳሽቦርድ ላይ በሚታየው መግለጫ ብቻ አይወሰኑ፤ የጉዳዩን ዝርዝር ይክፈቱ። የሁኔታ ታሪክ፣ ችሎቶች፣ ውሳኔዎች፣ ፋይሎችና የድርጊት ጥያቄዎችን ይመልከቱ።", "የሁኔታ ለውጥ በሥርዓቱ ላይ ሲታይ ወይም በተፈቀደ የተቋም ማስታወቂያ ሲወጣ ብቻ ተቀባይነት አለው።", "ስህተት ካዩ የጉዳይ መዝገብ ቁጥር፣ የማያ ገጽ ምስል፣ ቀን/ሰዓትና አጭር መግለጫ አካትተው ሪፖርት ያድርጉ።"]),
12: (["Use the case message thread for case-related communication only. State the case number, purpose, and requested action clearly.", "Do not use messages for emergencies or to bypass a formal filing, response, appeal, or evidence process.", "Check notifications regularly; marking a notification as read does not complete the requested action."], ["የጉዳዩን የመልዕክት ክፍል ከጉዳዩ ጋር ለተያያዘ ግንኙነት ብቻ ይጠቀሙ። የመዝገብ ቁጥሩን፣ ዓላማውንና የተጠየቀውን ድርጊት በግልጽ ይጻፉ።", "መልዕክትን ለአስቸኳይ ጉዳይ ወይም መደበኛ ማመልከቻ፣ መልስ፣ ይግባኝ ወይም ማስረጃ ሂደትን ለማለፍ አይጠቀሙ።", "ማስታወቂያዎችን በተደጋጋሚ ይፈትሹ፤ ማስታወቂያን እንደታየ ምልክት ማድረግ የተጠየቀውን ድርጊት አያጠናቅቅም።"]),
13: (["Confirm that the case is eligible and that the filing period has not expired. State the challenged decision, grounds, and requested relief precisely.", "Attach only the required records and verify submission. Keep the appeal reference and receipt.", "An appeal is governed by applicable procedure and authorization; system availability does not extend a legal deadline."], ["ጉዳዩ ለይግባኝ ብቁ መሆኑንና የማቅረቢያ ጊዜው አለማለፉን ያረጋግጡ። የተቃወሙትን ውሳኔ፣ ምክንያቶችና የተጠየቀውን ውጤት በግልጽ ይግለጹ።", "አስፈላጊ መዝገቦችን ብቻ ያያይዙ፤ ማስገባቱን ያረጋግጡ። የይግባኝ ማጣቀሻና ደረሰኝ ያስቀምጡ።", "ይግባኝ በተፈጻሚው ሥነ ሥርዓትና በተሰጠ ሥልጣን ይመራል፤ ሥርዓቱ መኖሩ ሕጋዊ የጊዜ ገደብን አያራዝምም።"]),
14: (["Keep identity, email, telephone, and address information current. Changes may be logged and may require verification.", "Role switching changes the working view, not ownership or permission. Complete pending work before switching.", "Sign out after use, particularly on a device not controlled by you."], ["የማንነት፣ የኢሜይል፣ የስልክና የአድራሻ መረጃዎን ወቅታዊ ያድርጉ። ለውጦች በኦዲት መዝገብ ሊመዘገቡና ማረጋገጫ ሊጠይቁ ይችላሉ።", "ሚና መቀየር የሥራ እይታን ብቻ ይቀይራል፤ ባለቤትነትን ወይም ፈቃድን አይቀይርም። ከመቀየርዎ በፊት የቀሩ ሥራዎችን ያጠናቅቁ።", "በተለይ በእርስዎ ቁጥጥር ሥር ባልሆነ መሣሪያ ላይ ሲጠቀሙ ከመለያዎ ይውጡ።"]),
15: (["The respondent account is personal. Register with accurate identity/contact data and verify the email before accessing protected case functions.", "If you already have an applicant account, use the available role-switch function rather than creating an unnecessary duplicate.", "Access remains limited to cases lawfully linked to the authenticated user."], ["የመልስ ሰጪ (ተጠሪ) መለያ የግል ነው። ትክክለኛ የማንነትና የመገኛ መረጃ ያስገቡ፤ የተጠበቁ የጉዳይ ተግባራትን ከመጠቀምዎ በፊት ኢሜይሉን ያረጋግጡ።", "የአመልካች መለያ ካለዎት አላስፈላጊ ተደጋጋሚ መለያ ከመፍጠር ይልቅ ያለውን የሚና መቀየሪያ ይጠቀሙ።", "መዳረሻ በሕጋዊ ሁኔታ ከተረጋገጠው ተጠቃሚ ጋር በተገናኙ ጉዳዮች ብቻ የተገደበ ነው።"]),
16: (["Enter the case number exactly as issued. Confirm the displayed parties and subject before linking or acting.", "Do not repeatedly guess case numbers. An unsuccessful or unauthorized search must be reported through support.", "Finding a case does not grant permission to view or respond; identity and case linkage controls still apply."], ["የጉዳይ መዝገብ ቁጥሩን እንደተሰጠው በትክክል ያስገቡ። ከማገናኘት ወይም ድርጊት ከመፈጸምዎ በፊት የታዩትን ተከራካሪ ወገኖችና ርዕስ ያረጋግጡ።", "የመዝገብ ቁጥሮችን በግምት በተደጋጋሚ አይሞክሩ። ያልተሳካ ወይም ያልተፈቀደ ፍለጋ ለድጋፍ ክፍል መገለጽ አለበት።", "ጉዳይን በፍለጋ ማግኘት የማየት ወይም መልስ የመስጠት ፈቃድ አይሰጥም፤ የማንነትና የጉዳይ ግንኙነት ቁጥጥር ተፈጻሚ ይሆናል።"]),
17: (["Read the applicant's filing and permitted attachments in full. Address each relevant allegation, state defenses clearly, and attach supporting records.", "Review names, dates, facts, declaration, and files before submission. Save the reference/receipt after submission.", "Submit within the displayed or officially notified deadline. Contact the institution for procedural guidance, not legal advice."], ["የአመልካቹን ማመልከቻና የተፈቀዱ አባሪዎችን ሙሉ በሙሉ ያንብቡ። ለእያንዳንዱ አግባብነት ያለው ክስ ምላሽ ይስጡ፣ መከላከያዎን በግልጽ ይግለጹ፣ ደጋፊ መዝገቦችን ያያይዙ።", "ከማስገባትዎ በፊት ስሞችን፣ ቀኖችን፣ ፍሬ ነገሮችን፣ መግለጫውንና ፋይሎችን ይገምግሙ። ካስገቡ በኋላ ማጣቀሻውን/ደረሰኙን ያስቀምጡ።", "በሥርዓቱ የታየው ወይም በይፋ የተገለጸው ጊዜ ገደብ ሳያልፍ ያስገቡ። ለሥነ ሥርዓት መመሪያ ተቋሙን ያነጋግሩ፤ የሕግ ምክር አይጠይቁ።"]),
18: (["Edit or withdraw only while the function is available and before the controlling deadline. Review the confirmation message carefully.", "Keep a copy of the final submitted version. A withdrawn response may no longer be considered unless lawfully resubmitted.", "Applicant replies are read-only unless a separate permitted action is displayed."], ["የማስተካከያ ወይም የማንሳት ተግባሩ በሚገኝበትና የጊዜ ገደቡ ሳያልፍ ብቻ ይፈጽሙ። የማረጋገጫ መልዕክቱን በጥንቃቄ ያንብቡ።", "የመጨረሻውን የገባ ቅጂ ያስቀምጡ። የተነሳ መልስ በሕጋዊ ሁኔታ እንደገና ካልቀረበ በስተቀር ግምት ውስጥ ላይገባ ይችላል።", "የአመልካች ምላሾች የተለየ የተፈቀደ ተግባር ካልታየ በስተቀር ለንባብ ብቻ ናቸው።"]),
19: (["Review notifications, case events, and deadlines after every sign-in. Open the related case to confirm context.", "Role switching changes the portal context. Verify the active role before submitting any action.", "Email or Telegram alerts are supplementary; the authenticated case record is the controlling source."], ["በእያንዳንዱ መግቢያ ጊዜ ማስታወቂያዎችን፣ የጉዳይ ክስተቶችንና የጊዜ ገደቦችን ይገምግሙ። ዐውዱን ለማረጋገጥ ተያያዥ ጉዳዩን ይክፈቱ።", "ሚና መቀየር የመግቢያውን የሥራ ዐውድ ይቀይራል። ማንኛውንም ድርጊት ከማስገባትዎ በፊት ንቁውን ሚና ያረጋግጡ።", "የኢሜይል ወይም የTelegram ማሳወቂያዎች ተጨማሪ ናቸው፤ በተረጋገጠው መለያ ውስጥ ያለው የጉዳይ መዝገብ ዋና ምንጭ ነው።"]),
20: (["Staff access is based on assigned role, permission, team, and case scope. Never use another officer's account.", "At sign-in, verify the environment, institution identity, active role, dashboard alerts, and pending work.", "All material actions may be audited. Sign out or lock the workstation whenever unattended."], ["የሥራ ኃላፊ መዳረሻ በተመደበ ሚና፣ ፈቃድ፣ ቡድንና የጉዳይ ወሰን ይወሰናል። የሌላ የሥራ ኃላፊ መለያ ፈጽሞ አይጠቀሙ።", "ሲገቡ የሥራ አካባቢውን፣ የተቋሙን መለያ፣ ንቁ ሚናውን፣ የዳሽቦርድ ማሳወቂያዎችንና ያልተጠናቀቁ ሥራዎችን ያረጋግጡ።", "ጉልህ የሆኑ ድርጊቶች በኦዲት መዝገብ ሊመዘገቡ ይችላሉ። የሥራ ጣቢያውን ለቀው ሲሄዱ መለያውን ይዝጉ ወይም መሣሪያውን ይቆልፉ።"]),
21: (["Check completeness, identity, jurisdiction, duplication, required files, and legibility. Record findings in the designated review fields.", "Accept, return for correction, or reject only under the assigned authority and applicable institutional procedure.", "Every decision must have a clear reason. Do not alter the applicant's original submission outside an authorized correction process."], ["ሙሉነትን፣ ማንነትን፣ ሥልጣንን፣ ተደጋጋሚነትን፣ አስፈላጊ ፋይሎችንና ተነባቢነትን ይፈትሹ። ግኝቶችን በተዘጋጀው የግምገማ መስክ ይመዝግቡ።", "ተቀባይነት መስጠት፣ ለማስተካከያ መመለስ ወይም ውድቅ ማድረግ በተሰጠ ሥልጣንና በተፈጻሚ የተቋም ሥነ ሥርዓት መሠረት ብቻ ይፈጸም።", "እያንዳንዱ ውሳኔ ግልጽ ምክንያት ሊኖረው ይገባል። ከተፈቀደ የማስተካከያ ሂደት ውጭ የአመልካቹን ዋና ማስገቢያ አይቀይሩ።"]),
22: (["Confirm acceptance, case type, workload, team scope, competence, and absence of conflict before assignment.", "Assign to the authorized officer/team, add a concise note where required, and verify that the new assignee is displayed.", "Reassignment must be justified and traceable; it must not be used to bypass workload or segregation controls."], ["ከመመደብዎ በፊት ተቀባይነትን፣ የጉዳይ ዓይነትን፣ የሥራ ጫናን፣ የቡድን ወሰንን፣ ብቃትንና የጥቅም ግጭት አለመኖሩን ያረጋግጡ።", "ለተፈቀደለት የሥራ ኃላፊ/ቡድን ይመድቡ፣ ካስፈለገ አጭር ማስታወሻ ይጨምሩ፣ አዲሱ ተመዳቢ መታየቱን ያረጋግጡ።", "ዳግም ምደባ ምክንያት ያለውና ክትትል የሚቻልበት መሆን አለበት፤ የሥራ ጫና ወይም የኃላፊነት መለያየት ቁጥጥርን ለማለፍ አይጠቀሙበት።"]),
23: (["Check party availability, statutory time requirements, assigned room/location, and institutional calendar before saving.", "Record date, time, place, purpose, presiding officer where applicable, and a clear note. Verify the generated notice/calendar file.", "For adjournment or cancellation, record the reason and ensure affected parties receive the approved notification."], ["ከማስቀመጥዎ በፊት የተከራካሪ ወገኖችን አቅርቦት፣ ሕጋዊ የጊዜ መስፈርት፣ የተመደበ ክፍል/ቦታና የተቋሙን የቀን መቁጠሪያ ይፈትሹ።", "ቀን፣ ሰዓት፣ ቦታ፣ ዓላማ፣ አግባብ ሲኖረው የሚመራውን ባለሥልጣንና ግልጽ ማስታወሻ ይመዝግቡ። የተፈጠረውን ማስታወቂያ/የቀን መቁጠሪያ ፋይል ያረጋግጡ።", "ለቀጠሮ ማራዘሚያ ወይም ስረዛ ምክንያቱን ይመዝግቡ፤ ተጽዕኖ ያለባቸው ወገኖች የተፈቀደውን ማሳወቂያ መቀበላቸውን ያረጋግጡ።"]),
24: (["Draft from the complete record. Select the correct decision type/template and state issues, findings, reasoning, outcome, and operative order consistently.", "Save versions, proofread names/citations/dates, and submit through the required review and approval chain.", "A draft is not final. Publish or communicate only the approved version and preserve the approved record."], ["ረቂቁን ከተሟላው የጉዳይ መዝገብ ያዘጋጁ። ትክክለኛውን የውሳኔ ዓይነት/ንድፍ ይምረጡ፤ ጭብጦችን፣ ግኝቶችን፣ ምክንያቶችን፣ ውጤትንና ተፈጻሚ ትእዛዝን በተጣጣመ ሁኔታ ይግለጹ።", "ቅጂዎችን ያስቀምጡ፣ ስሞችን/ዋቢዎችን/ቀኖችን ያርሙ፣ በተወሰነው የግምገማና የማጽደቅ ሰንሰለት ያስተላልፉ።", "ረቂቅ የመጨረሻ ውሳኔ አይደለም። የጸደቀውን ቅጂ ብቻ ያትሙ ወይም ያሳውቁ፤ የጸደቀውን መዝገብ ይጠብቁ።"]),
25: (["Bench notes are restricted working records. Write factual, professional notes and associate them with the correct case/hearing.", "Do not place a final order, public communication, or evidence alteration in a private note.", "Access, editing, and deletion are limited by role and audit controls."], ["የዳኛ ማስታወሻዎች የተገደቡ የሥራ መዝገቦች ናቸው። ተጨባጭና ሙያዊ ማስታወሻ ይጻፉ፤ ከትክክለኛው ጉዳይ/ችሎት ጋር ያገናኙ።", "የመጨረሻ ትእዛዝ፣ ለሕዝብ የሚገለጽ ግንኙነት ወይም የማስረጃ ለውጥ በግል ማስታወሻ ውስጥ አያስቀምጡ።", "መዳረሻ፣ ማስተካከያና ስረዛ በሚናና በኦዲት ቁጥጥር የተገደቡ ናቸው።"]),
26: (["Choose the correct approved template/category and recipient. Verify case linkage, reference number, subject, names, address, body, attachments, and signatory.", "Submit through the required approval sequence. Record approval/rejection reasons and correct the draft without overwriting an approved record.", "Preview the final public copy before issue and retain the institutional copy according to records policy."], ["ትክክለኛውን የጸደቀ ንድፍ/ምድብና ተቀባይ ይምረጡ። የጉዳይ ግንኙነትን፣ የማጣቀሻ ቁጥርን፣ ርዕስን፣ ስሞችን፣ አድራሻን፣ የደብዳቤ ይዘትን፣ አባሪዎችንና ፈራሚውን ያረጋግጡ።", "በተወሰነው የማጽደቅ ሰንሰለት ያስተላልፉ። የማጽደቅ/ውድቅ ምክንያትን ይመዝግቡ፤ የጸደቀ መዝገብ ሳይተካ ረቂቁን ያስተካክሉ።", "ከመላክዎ በፊት የመጨረሻውን የሕዝብ ቅጂ ይመልከቱ፤ የተቋሙን ቅጂ በመዝገብ አያያዝ ፖሊሲ መሠረት ያስቀምጡ።"]),
27: (["Create an inspection request only within assigned authority. Define subject, scope, date, responsible inspector, due date, and required records.", "The inspector records objective findings, evidence references, recommendation, and completion status. Supervisory review follows where configured.", "Maintain independence, confidentiality, and chain of custody for inspection material."], ["የፍተሻ ጥያቄን በተሰጠ ሥልጣን ውስጥ ብቻ ይፍጠሩ። ርዕስ፣ ወሰን፣ ቀን፣ ኃላፊ መርማሪ፣ የማጠናቀቂያ ጊዜና አስፈላጊ መዝገቦችን ይወስኑ።", "መርማሪው ተጨባጭ ግኝቶችን፣ የማስረጃ ዋቢዎችን፣ ምክረ ሐሳብንና የማጠናቀቂያ ሁኔታን ይመዘግባል። በሥርዓቱ ከተዋቀረ የኃላፊ ግምገማ ይከተላል።", "የፍተሻ ገለልተኝነትን፣ ምስጢራዊነትንና የማስረጃ አያያዝ ሰንሰለትን ይጠብቁ።"]),
28: (["Verify filing authority, timeliness, challenged decision, grounds, documents, and fee/waiver information where applicable.", "Track review, assignment, hearing, outcome, and linked records without changing the source-case history.", "Only authorized officers may accept, decide, close, or publish appeal information."], ["የማቅረብ ሥልጣንን፣ ወቅታዊነትን፣ የተቃውሞ ውሳኔን፣ ምክንያቶችን፣ ሰነዶችንና አግባብ ሲኖረው የክፍያ/ነጻ መሆን መረጃን ያረጋግጡ።", "የምንጭ ጉዳዩን ታሪክ ሳይቀይሩ ግምገማን፣ ምደባን፣ ችሎትን፣ ውጤትንና ተያያዥ መዝገቦችን ይከታተሉ።", "የይግባኝ መረጃን መቀበል፣ መወሰን፣ መዝጋት ወይም ማተም የሚችሉት የተፈቀደላቸው የሥራ ኃላፊዎች ብቻ ናቸው።"]),
29: (["Select the correct officer, evaluation period, criteria, and authorized evaluator. Base ratings on verifiable institutional records.", "Add balanced evidence and comments, complete required review/acknowledgment, and lock or finalize only when authorized.", "Performance information is confidential personnel data and must not be exported or shared without authority."], ["ትክክለኛውን የሥራ ኃላፊ፣ የግምገማ ጊዜ፣ መስፈርቶችና የተፈቀደ ገምጋሚ ይምረጡ። ደረጃ አሰጣጡን ሊረጋገጥ በሚችል የተቋም መዝገብ ላይ ይመሥርቱ።", "ሚዛናዊ ማስረጃና አስተያየት ይጨምሩ፣ አስፈላጊውን ግምገማ/እውቅና ያጠናቅቁ፣ ሥልጣን ሲኖር ብቻ ያጽድቁ ወይም ይቆልፉ።", "የአፈጻጸም መረጃ ምስጢራዊ የሠራተኛ መረጃ ነው፤ ያለ ሥልጣን ወደ ውጭ አይውጣ ወይም አይጋራ።"]),
30: (["Apply filters deliberately and verify date range, status, team, case type, and scope before relying on totals.", "For an official report, record report name, parameters, generation date/time, responsible officer, and confidentiality classification.", "Exports may contain protected data. Store, transmit, print, and dispose of them under institutional records and security rules."], ["ማጣሪያዎችን በጥንቃቄ ይተግብሩ፤ በድምር ውጤት ከመመሥረትዎ በፊት የቀን ክልል፣ ሁኔታ፣ ቡድን፣ የጉዳይ ዓይነትና ወሰን ያረጋግጡ።", "ለይፋዊ ሪፖርት የሪፖርት ስም፣ መለኪያዎች፣ የተፈጠረበት ቀን/ሰዓት፣ ኃላፊ ባለሥልጣንና የምስጢራዊነት ደረጃ ይመዝግቡ።", "ወደ ውጭ የወጡ ፋይሎች የተጠበቀ መረጃ ሊኖራቸው ይችላል። በተቋሙ የመዝገብና የደህንነት ደንብ መሠረት ያስቀምጡ፣ ያስተላልፉ፣ ያትሙ ወይም ያስወግዱ።"]),
31: (["Create a staff account only from an approved request. Verify official identity, email, team, employment status, role, and effective date.", "Apply least privilege. Separate user activation from role/permission approval when institutional procedure requires it.", "Disable access promptly on transfer, suspension, or separation; retain audit history rather than deleting accountable records."], ["የሥራ ኃላፊ መለያን በጸደቀ ጥያቄ ብቻ ይፍጠሩ። ይፋዊ ማንነትን፣ ኢሜይልን፣ ቡድንን፣ የቅጥር ሁኔታን፣ ሚናንና የሚጀምርበትን ቀን ያረጋግጡ።", "ዝቅተኛውን አስፈላጊ ፈቃድ ብቻ ይስጡ። የተቋሙ ሥነ ሥርዓት ሲጠይቅ መለያ ማንቃትን ከሚና/ፈቃድ ማጽደቅ ይለዩ።", "ሲዛወር፣ ሲታገድ ወይም ከሥራ ሲለቅ መዳረሻውን ወዲያውኑ ያሰናክሉ፤ ተጠያቂነት ያላቸውን መዝገቦች ከመሰረዝ ይልቅ የኦዲት ታሪኩን ይጠብቁ።"]),
32: (["Change system identity, locale, notification, content, backup, and operational options only through approved change control.", "Review the current value, record the request/approval, change one controlled item at a time, verify the result, and retain evidence.", "CMS content, terms, announcements, templates, and permissions require separate accuracy and publication review. Audit logs are read-only evidence."], ["የሥርዓት መለያ፣ ቋንቋ፣ ማሳወቂያ፣ ይዘት፣ ምትኬና የሥራ ክንውን ቅንብሮችን በጸደቀ የለውጥ ቁጥጥር ብቻ ይቀይሩ።", "ያለውን እሴት ይገምግሙ፣ ጥያቄውን/ማጽደቁን ይመዝግቡ፣ በአንድ ጊዜ አንድ ቁጥጥር ያለበት ንጥል ይቀይሩ፣ ውጤቱን ያረጋግጡ፣ ማስረጃውን ይያዙ።", "የCMS ይዘት፣ ደንቦች፣ ማስታወቂያዎች፣ ንድፎችና ፈቃዶች የተለየ የትክክለኛነትና የህትመት ግምገማ ይፈልጋሉ። የኦዲት መዝገቦች ለንባብ ብቻ የሚጠበቁ ማስረጃዎች ናቸው።"]),
33: (["Changing language changes labels and displayed content; it does not change case data, permissions, deadlines, or the legal meaning of records.", "If a translation appears unclear, retain the case reference and report the exact page/label to authorized support."], ["ቋንቋ መቀየር መግለጫዎችንና የሚታየውን ይዘት ይቀይራል፤ የጉዳይ መረጃን፣ ፈቃዶችን፣ የጊዜ ገደቦችን ወይም የመዝገቦችን ሕጋዊ ትርጉም አይቀይርም።", "ትርጉም ግልጽ ካልሆነ የጉዳዩን ማጣቀሻ ይያዙ፤ ትክክለኛውን ገጽ/መግለጫ ለተፈቀደ ድጋፍ ያሳውቁ።"]),
34: (["Use a long, unique password that is not used for email or another service. Never place it in a document, browser note, or shared message.", "If compromise is suspected, change the password immediately, sign out of other sessions where available, and notify the designated security contact."], ["ረጅምና ልዩ የሆነ፣ ለኢሜይል ወይም ለሌላ አገልግሎት ያልተጠቀሙበት የይለፍ ቃል ይጠቀሙ። በሰነድ፣ በድር አሳሽ ማስታወሻ ወይም በጋራ መልዕክት ውስጥ አያስቀምጡት።", "መለያው ተጋልጧል ብለው ከጠረጠሩ የይለፍ ቃሉን ወዲያውኑ ይቀይሩ፣ ካለ ከሌሎች ክፍት ጊዜያት ይውጡ፣ ለተመደበው የደህንነት ኃላፊ ያሳውቁ።"]),
35: (["Before reporting a problem, refresh once, verify internet connection, browser, active role, and required fields. Do not repeatedly submit the same action.", "A useful support report includes: user role, page, case/reference number, exact error, date/time, steps taken, and a screenshot with unrelated personal data hidden.", "Never send passwords, OTPs, full identity documents, or confidential case files through an unapproved support channel."], ["ችግርን ከማሳወቅዎ በፊት ገጹን አንድ ጊዜ ያድሱ፣ ኢንተርኔትን፣ የድር አሳሹን፣ ንቁ ሚናውንና ግዴታ መስኮችን ያረጋግጡ። ተመሳሳይ ድርጊትን በተደጋጋሚ አያስገቡ።", "ጠቃሚ የድጋፍ ሪፖርት የተጠቃሚ ሚና፣ ገጽ፣ የጉዳይ/ማጣቀሻ ቁጥር፣ ትክክለኛ ስህተት፣ ቀን/ሰዓት፣ የተከናወኑ ደረጃዎችና አግባብ የሌለው የግል መረጃ የተሸፈነበት የማያ ገጽ ምስል ያካትታል።", "የይለፍ ቃል፣ OTP፣ ሙሉ የማንነት ሰነድ ወይም ምስጢራዊ የጉዳይ ፋይል ባልተፈቀደ የድጋፍ መንገድ አይላኩ።"]),
36: (["Administrative terminology: 'case record' means the controlled electronic collection for one matter; 'audit log' means the trace of significant system actions; 'scope' means the records and functions an authenticated user is authorized to access."], ["የአስተዳደር ቃላት፡ 'የጉዳይ መዝገብ' ማለት ለአንድ ጉዳይ በቁጥጥር ሥር የተያዘ የኤሌክትሮኒክ ሰነዶች ስብስብ ነው፤ 'የኦዲት መዝገብ' ማለት ጉልህ የሥርዓት ድርጊቶችን የሚከታተል ማስረጃ ነው፤ 'ወሰን' ማለት የተረጋገጠ ተጠቃሚ እንዲደርስባቸው የተፈቀዱ መዝገቦችና ተግባራት ናቸው።"]),
}


def add_detail_row(table, en, am):
    row = table.add_row()
    set_cell_text(row.cells[0], en, amharic=False, bold_first=True)
    set_cell_text(row.cells[1], am, amharic=True, bold_first=True)
    shade(row.cells[0], PALE)
    shade(row.cells[1], PALE)


def add_bilingual_heading(doc, en, am, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{en}  |  {am}")
    set_run_font(r, "Arial", 15 if level == 1 else 12, True, BLUE if level == 1 else MID_BLUE)
    return p


def add_bilingual_table(doc, rows):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(3.55)
    t.columns[1].width = Inches(3.55)
    set_cell_text(t.rows[0].cells[0], ["English"], bold_first=True)
    set_cell_text(t.rows[0].cells[1], ["አማርኛ"], amharic=True, bold_first=True)
    for c in t.rows[0].cells:
        shade(c, BLUE)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
    for en, am in rows:
        row = t.add_row()
        set_cell_text(row.cells[0], en if isinstance(en, list) else [en])
        set_cell_text(row.cells[1], am if isinstance(am, list) else [am], amharic=True)
        if len(t.rows) % 2 == 0:
            shade(row.cells[0], PALE)
            shade(row.cells[1], PALE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def add_appendices(doc):
    doc.add_page_break()
    add_bilingual_heading(doc, "6. Institutional Operating and Security Controls", "6. የተቋም የሥራ ክንውንና የደህንነት ቁጥጥሮች")
    add_bilingual_table(doc, [
        (["Authorization and accountability", "- Use only the role, team, case scope, and permissions officially assigned to you.", "- Never approve your own work where maker-checker separation applies.", "- Material actions are attributable to the authenticated account and may be audited."], ["ፈቃድና ተጠያቂነት", "- በይፋ የተመደበልዎትን ሚና፣ ቡድን፣ የጉዳይ ወሰንና ፈቃድ ብቻ ይጠቀሙ።", "- አዘጋጅና አጽዳቂ እንዲለያዩ በሚጠየቅበት ሂደት የራስዎን ሥራ አያጽድቁ።", "- ጉልህ ድርጊቶች ከተረጋገጠው መለያ ጋር ይያያዛሉ እና በኦዲት ሊመረመሩ ይችላሉ።"]),
        (["Information protection", "- Collect and expose only information required for the authorized task.", "- Verify the recipient before sending a letter, export, notice, or case file.", "- Store downloaded records only in approved institutional locations and remove temporary copies securely."], ["የመረጃ ጥበቃ", "- ለተፈቀደው ተግባር የሚያስፈልግ መረጃ ብቻ ይሰብስቡ እና ያሳዩ።", "- ደብዳቤ፣ ወደ ውጭ የወጣ ፋይል፣ ማስታወቂያ ወይም የጉዳይ መዝገብ ከመላክዎ በፊት ተቀባዩን ያረጋግጡ።", "- የወረዱ መዝገቦችን በተፈቀደ የተቋም ማከማቻ ብቻ ያስቀምጡ፤ ጊዜያዊ ቅጂዎችን በደህና ያስወግዱ።"]),
        (["Record integrity and continuity", "- Do not overwrite an approved decision, letter, inspection finding, or closed record.", "- Use the designated correction, version, reopening, or reassignment function and document the reason.", "- Follow the institution's backup, retention, restoration, and continuity procedure."], ["የመዝገብ ታማኝነትና ቀጣይነት", "- የጸደቀ ውሳኔ፣ ደብዳቤ፣ የፍተሻ ግኝት ወይም የተዘጋ መዝገብ አይተኩ።", "- የተዘጋጀውን የማስተካከያ፣ የቅጂ፣ ዳግም መክፈቻ ወይም ዳግም ምደባ ተግባር ይጠቀሙ፤ ምክንያቱን ይመዝግቡ።", "- የተቋሙን የምትኬ፣ የማቆያ፣ የመልሶ ማቋቋምና የሥራ ቀጣይነት ሥነ ሥርዓት ይከተሉ።"]),
    ])

    add_bilingual_heading(doc, "7. Role and Responsibility Matrix", "7. የሚናና የኃላፊነት ማትሪክስ")
    add_bilingual_table(doc, [
        (["Applicant", "Creates and verifies an account; files and tracks own cases; supplies documents/witnesses; receives notices; replies and appeals when eligible."], ["አመልካች", "መለያ ይፈጥራልና ያረጋግጣል፤ የራሱን ጉዳይ ያስመዘግባልና ይከታተላል፤ ሰነዶችን/ምስክሮችን ያቀርባል፤ ማስታወቂያ ይቀበላል፤ ብቁ ሲሆን ምላሽና ይግባኝ ያቀርባል።"]),
        (["Respondent", "Finds a lawfully linked case; reviews the filing; submits, edits, or withdraws a response within the permitted period; tracks replies and notices."], ["መልስ ሰጪ (ተጠሪ)", "በሕጋዊ ሁኔታ የተገናኘውን ጉዳይ ያገኛል፤ ማመልከቻውን ይገመግማል፤ በተፈቀደው ጊዜ መልስ ያስገባል፣ ያስተካክላል ወይም ያነሳል፤ ምላሾችንና ማስታወቂያዎችን ይከታተላል።"]),
        (["Registrar / intake officer", "Checks completeness and jurisdiction; records review findings; accepts, returns, or rejects under authority; preserves the original submission."], ["መዝገብ ያዥ / የመቀበያ የሥራ ኃላፊ", "ሙሉነትንና ሥልጣንን ይፈትሻል፤ የግምገማ ግኝትን ይመዘግባል፤ በሥልጣኑ መሠረት ይቀበላል፣ ለማስተካከያ ይመልሳል ወይም ውድቅ ያደርጋል፤ ዋናውን ማስገቢያ ይጠብቃል።"]),
        (["Assigned officer / adjudicator", "Manages hearings and case progress; records professional notes; prepares decisions and related records within assigned authority."], ["ተመዳቢ የሥራ ኃላፊ / ውሳኔ ሰጪ", "ችሎቶችንና የጉዳይ ሂደትን ያስተዳድራል፤ ሙያዊ ማስታወሻ ይመዘግባል፤ በተሰጠ ሥልጣን ውስጥ ውሳኔና ተያያዥ መዝገቦችን ያዘጋጃል።"]),
        (["Approver / supervisor", "Reviews accuracy, authority, completeness, and segregation of duties; approves or returns work with a documented reason."], ["አጽዳቂ / ኃላፊ", "ትክክለኛነትን፣ ሥልጣንን፣ ሙሉነትንና የኃላፊነት መለያየትን ይገመግማል፤ ሥራውን ያጸድቃል ወይም ምክንያቱን በመመዝገብ ይመልሳል።"]),
        (["System administrator", "Manages approved users, roles, permissions, settings, content, and operational controls; does not decide case merits by virtue of technical access."], ["የሥርዓት አስተዳዳሪ", "የጸደቁ ተጠቃሚዎችን፣ ሚናዎችን፣ ፈቃዶችን፣ ቅንብሮችን፣ ይዘትንና የሥራ ክንውን ቁጥጥሮችን ያስተዳድራል፤ በቴክኒክ መዳረሻ ብቻ የጉዳይን ፍሬ ነገር አይወስንም።"]),
    ])

    add_bilingual_heading(doc, "8. Incident and Support Escalation", "8. የክስተትና የድጋፍ ከፍ ማድረጊያ ሂደት")
    add_bilingual_table(doc, [
        (["1. Protect", "Stop the affected action. Do not retry destructive or duplicate submissions. Preserve the screen and reference information."], ["1. ይጠብቁ", "ችግሩ ያጋጠመውን ድርጊት ያቁሙ። አጥፊ ወይም ተደጋጋሚ ማስገቢያን እንደገና አይሞክሩ። የማያ ገጹንና ማጣቀሻ መረጃውን ይጠብቁ።"]),
        (["2. Record", "Note date/time, user role, environment, page, case/reference, exact message, action attempted, and business impact."], ["2. ይመዝግቡ", "ቀን/ሰዓት፣ የተጠቃሚ ሚና፣ የሥራ አካባቢ፣ ገጽ፣ የጉዳይ/ማጣቀሻ ቁጥር፣ ትክክለኛ መልዕክት፣ የተሞከረ ድርጊትና በሥራ ላይ ያሳደረውን ተጽዕኖ ይመዝግቡ።"]),
        (["3. Report", "Use the institution's approved service desk or supervisory channel. Classify suspected unauthorized access, data exposure, or account compromise as a security incident."], ["3. ያሳውቁ", "የተቋሙን የተፈቀደ የድጋፍ ማዕከል ወይም የኃላፊ መንገድ ይጠቀሙ። ያልተፈቀደ መዳረሻ፣ የመረጃ መጋለጥ ወይም የመለያ ጥሰት ጥርጣሬን እንደ የደህንነት ክስተት ይመድቡ።"]),
        (["4. Verify closure", "After resolution, confirm the expected record state, attach resolution evidence to the ticket where required, and close only with user/owner confirmation."], ["4. መዘጋቱን ያረጋግጡ", "ችግሩ ከተፈታ በኋላ የሚጠበቀውን የመዝገብ ሁኔታ ያረጋግጡ፤ ካስፈለገ የመፍትሔ ማስረጃውን ከትኬቱ ጋር ያያይዙ፤ በተጠቃሚ/ባለቤት ማረጋገጫ ብቻ ይዝጉ።"]),
    ])

    add_bilingual_heading(doc, "9. Daily and Periodic Checklists", "9. ዕለታዊና ወቅታዊ የማረጋገጫ ዝርዝሮች")
    add_bilingual_table(doc, [
        (["Daily - all users", "- Verify the active role and unread notifications.\n- Review deadlines and pending actions.\n- Save receipts/references.\n- Sign out when finished."], ["ዕለታዊ - ለሁሉም ተጠቃሚዎች", "- ንቁ ሚናውንና ያልታዩ ማስታወቂያዎችን ያረጋግጡ።\n- የጊዜ ገደቦችንና ያልተጠናቀቁ ድርጊቶችን ይገምግሙ።\n- ደረሰኞችን/ማጣቀሻዎችን ያስቀምጡ።\n- ሲጨርሱ ከመለያዎ ይውጡ።"]),
        (["Daily - staff", "- Review intake, assigned cases, hearings, approvals, and escalations.\n- Resolve or document overdue work.\n- Confirm that issued records use approved versions."], ["ዕለታዊ - ለሥራ ኃላፊዎች", "- የመቀበያ ወረፋ፣ የተመደቡ ጉዳዮች፣ ችሎቶች፣ ማጽደቆችና ከፍ ያሉ ጥያቄዎችን ይገምግሙ።\n- ጊዜያቸው ያለፈ ሥራዎችን ይፍቱ ወይም ምክንያቱን ይመዝግቡ።\n- የወጡ መዝገቦች የጸደቁ ቅጂዎችን መጠቀማቸውን ያረጋግጡ።"]),
        (["Monthly - administration", "- Review active users, transfers, leavers, roles, and privileged access.\n- Review backup evidence, audit exceptions, failed notifications, and unresolved incidents.\n- Confirm content, templates, and institutional contact details remain current."], ["ወርሃዊ - ለአስተዳደር", "- ንቁ ተጠቃሚዎችን፣ ዝውውሮችን፣ ከሥራ የለቀቁትን፣ ሚናዎችንና ልዩ መዳረሻዎችን ይገምግሙ።\n- የምትኬ ማስረጃ፣ የኦዲት ልዩነቶች፣ ያልተሳኩ ማሳወቂያዎችና ያልተፈቱ ክስተቶችን ይገምግሙ።\n- ይዘት፣ ንድፎችና የተቋሙ የመገኛ መረጃ ወቅታዊ መሆናቸውን ያረጋግጡ።"]),
    ])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("End of controlled user manual  |  የቁጥጥር ሥር ያለው የተጠቃሚ መመሪያ መጨረሻ")
    set_run_font(r, "Arial", 9, True, BLUE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(SOURCE)

    # Controlled-document metadata.
    control = doc.tables[1]
    control.cell(1, 1).text = "1.1"
    control.cell(2, 1).text = "17 July 2026"
    control.cell(3, 1).text = "Approved for institutional use - detailed bilingual edition"
    control.cell(4, 1).text = "System Owner / Authorized Government Institution"

    cover = doc.tables[0].cell(0, 0)
    for p in cover.paragraphs:
        for run in p.runs:
            old_text = run.text
            new_text = old_text.replace("የመዝገብ አስተዳደር ስርዓት", "የፍርድ ቤት ጉዳይ መዝገብ አስተዳደር ሥርዓት")
            new_text = new_text.replace("Bilingual —", "Detailed Bilingual -")
            if new_text != old_text:
                run.text = new_text

    # Normalize selected formal-government terminology throughout existing text.
    replacements = {
        "ስርዓት": "ሥርዓት",
        "የፍርድ ቤት ሰራተኞች": "የተቋሙ የሥራ ኃላፊዎች",
        "ተከሳሾች": "መልስ ሰጪዎች (ተጠሪዎች)",
        "የተከሳሽ መግቢያ": "የመልስ ሰጪ (ተጠሪ) መግቢያ",
        "ሰራተኛ": "የሥራ ኃላፊ",
        "ስልጣን": "ሥልጣን",
        "ስማት": "ችሎት",
    }
    for p in doc.paragraphs:
        for run in p.runs:
            old_text = run.text
            new_text = old_text
            for old, new in replacements.items():
                new_text = new_text.replace(old, new)
            if new_text != old_text:
                run.text = new_text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        old_text = run.text
                        new_text = old_text
                        for old, new in replacements.items():
                            new_text = new_text.replace(old, new)
                        if new_text != old_text:
                            run.text = new_text

    for idx, (en, am) in DETAILS.items():
        add_detail_row(doc.tables[idx], en, am)

    add_appendices(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    for sec in doc.sections:
        footer = sec.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = "CCMS User Manual v1.1 | Controlled Copy | የCCMS ተጠቃሚ መመሪያ ስሪት 1.1 | ቁጥጥር ያለበት ቅጂ"
        for run in p.runs:
            set_run_font(run, "Arial", 7.5, False, "666666")

    # Keep tables readable and prevent row splitting when possible.
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
            for cell in row.cells:
                margins(cell)

    props = doc.core_properties
    props.title = "CCMS Detailed Bilingual User Manual - English and Amharic"
    props.subject = "Institutional operating manual for the Court Case Management System"
    props.comments = "Revised detailed edition using formal government-organization Amharic terminology."
    props.version = "1.1"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
